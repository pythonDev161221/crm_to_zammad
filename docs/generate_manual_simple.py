"""
Run: cd /home/developer/crm_to_zammad/backend && ~/.local/bin/pipenv run python ../docs/generate_manual_simple.py
Output: /home/developer/crm_to_zammad/docs/manual_simple_ru.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def add_step(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Inches(0.3)
    p.add_run(text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3)
    p.add_run(text)
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run('Примечание: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0x8B, 0x44, 0x00)
    run2 = p.add_run(text)
    run2.italic = True
    run2.font.color.rgb = RGBColor(0x8B, 0x44, 0x00)
    return p


def add_screenshot_placeholder(doc, description):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(f'[СКРИНШОТ: {description}]')
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
    return p


# ─────────────────────────────────────────────
doc = Document()

style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)

# ══════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Система заявок технической поддержки')
run.bold = True
run.font.size = Pt(20)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Краткое руководство')
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_paragraph()
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('Для сотрудников и руководителей станций')
run3.italic = True
run3.font.size = Pt(13)

doc.add_page_break()

# ══════════════════════════════════════════════
# PART 1 — WORKER: REGISTRATION
# ══════════════════════════════════════════════
add_heading(doc, '1. Для сотрудника: регистрация по токену', level=1)
add_para(doc,
    'Перед началом работы вам нужно зарегистрироваться. '
    'Для этого руководитель станции пришлёт вам пригласительную ссылку или токен.')

doc.add_paragraph()

add_heading(doc, '1.1 Если руководитель прислал ссылку', level=2)
add_step(doc, 'Нажмите на ссылку в Telegram — откроется приложение.')
add_step(doc, 'Токен заполнится автоматически.')
add_step(doc, 'Введите своё имя в поле «Имя *».')
add_step(doc, 'При желании введите фамилию в поле «Фамилия».')
add_step(doc, 'Нажмите «Создать аккаунт».')

add_screenshot_placeholder(doc, 'Форма регистрации с заполненным полем «Токен приглашения *», полями «Имя *» и «Фамилия», кнопка «Создать аккаунт»')

doc.add_paragraph()

add_heading(doc, '1.2 Если у вас только токен (без ссылки)', level=2)
add_step(doc, 'Откройте приложение в Telegram.')
add_step(doc, 'На экране приветствия нажмите «Зарегистрироваться по приглашению».')
add_step(doc, 'Заполните три поля:')
add_bullet(doc, '«Токен приглашения *» — вставьте токен, который прислал руководитель.')
add_bullet(doc, '«Имя *» — ваше имя (обязательно).')
add_bullet(doc, '«Фамилия» — ваша фамилия (необязательно).')
add_step(doc, 'Нажмите «Создать аккаунт».')

add_screenshot_placeholder(doc, 'Форма регистрации: поля «Токен приглашения *», «Имя *», «Фамилия», кнопка «Создать аккаунт»')

add_note(doc,
    'Токен одноразовый — после вашей регистрации он становится недействительным. '
    'Если регистрация не прошла, попросите руководителя создать новый токен.')

doc.add_page_break()

# ══════════════════════════════════════════════
# PART 2 — WORKER: CREATE TICKET
# ══════════════════════════════════════════════
add_heading(doc, '2. Для сотрудника: создание заявки', level=1)
add_para(doc,
    'Чтобы сообщить о технической проблеме, создайте заявку. '
    'IT-специалист получит уведомление и возьмёт её в работу.')

doc.add_paragraph()

add_step(doc, 'На главном экране нажмите кнопку «+» в правом нижнем углу.')
add_step(doc, 'Введите краткое описание проблемы в поле «Тема».')
add_step(doc, 'При необходимости добавьте подробности в поле «Детали».')
add_step(doc, 'Если нужно — прикрепите фотографии, нажав «📷 Прикрепить фото».')
add_step(doc, 'Нажмите «Отправить».')

add_screenshot_placeholder(doc, 'Экран создания заявки: поля «Тема», «Детали», кнопка прикрепления фото, кнопка «Отправить»')

add_para(doc, 'После отправки заявка появится в списке. Возможные статусы:')
add_bullet(doc, '«Открыта» — заявка создана, ожидает IT-специалиста.')
add_bullet(doc, '«В работе» — IT-специалист принял заявку.')
add_bullet(doc, '«Решена» — проблема устранена.')

add_note(doc, 'Одна заявка — одна проблема. Если проблем несколько — создайте отдельные заявки.')

doc.add_page_break()

# ══════════════════════════════════════════════
# PART 3 — STATION MANAGER: INVITE
# ══════════════════════════════════════════════
add_heading(doc, '3. Для руководителя станции: приглашение сотрудника', level=1)
add_para(doc,
    'Чтобы добавить нового сотрудника на станцию, создайте пригласительную ссылку '
    'и отправьте её сотруднику.')

doc.add_paragraph()

add_heading(doc, '3.1 Как создать и отправить ссылку', level=2)
add_step(doc, 'На главном экране нажмите кнопку «Пригласить» в правом верхнем углу.')
add_step(doc, 'Откроется экран «Пригласительная ссылка».')
add_step(doc, 'Если ссылки нет — нажмите «Создать новую ссылку».')
add_step(doc, 'Нажмите «Скопировать ссылку».')
add_step(doc, 'Отправьте скопированную ссылку сотруднику в Telegram.')
add_step(doc, 'Сотрудник нажимает на ссылку, вводит имя и нажимает «Создать аккаунт».')

add_screenshot_placeholder(doc, 'Экран «Пригласительная ссылка»: текущая ссылка, кнопки «Скопировать ссылку», «Деактивировать ссылку», «Создать новую ссылку»')

doc.add_paragraph()

add_heading(doc, '3.2 Важные правила', level=2)
add_bullet(doc, 'Одна ссылка — для всех сотрудников станции. По ней могут зарегистрироваться несколько человек.')
add_bullet(doc, 'Ссылка остаётся активной до тех пор, пока вы её не деактивируете или не создадите новую.')
add_bullet(doc, 'Одновременно может существовать только одна активная ссылка на станцию.')
add_bullet(doc, 'Ненужную ссылку можно удалить кнопкой «Деактивировать ссылку».')

add_note(doc,
    'Если вы создадите новую ссылку — старая автоматически станет недействительной.')

# ══════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════
output_path = '/home/developer/crm_to_zammad/docs/manual_simple_ru.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
